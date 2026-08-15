"""Render the synthetic data room to PDF and DOCX.

Usage:
    python scripts/generate_corpus.py
    python scripts/generate_corpus.py --out data/synthetic
    python scripts/generate_corpus.py --verify   (check consistency only)
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from loupe.corpus.defects import primary_class
from loupe.corpus.registry import ALL_DOCUMENTS, CATEGORIES, PLANTED_DEFECTS, DocSpec


def write_pdf(spec: DocSpec, out_dir: Path) -> Path:
    """Render a document spec to PDF."""
    path = out_dir / spec.filename
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=A4, title=spec.title)
    flow: list[object] = [Paragraph(spec.title, styles["Heading1"]), Spacer(1, 12)]
    for para in spec.paragraphs:
        flow.append(Paragraph(para, styles["BodyText"]))
        flow.append(Spacer(1, 8))
    doc.build(flow)
    return path


def write_docx(spec: DocSpec, out_dir: Path) -> Path:
    """Render a document spec to DOCX."""
    path = out_dir / spec.filename
    docx = DocxDocument()
    docx.add_heading(spec.title, level=1)
    for para in spec.paragraphs:
        docx.add_paragraph(para)
    docx.save(str(path))
    return path


def render(spec: DocSpec, out_dir: Path) -> Path:
    """Render one document in the format its extension implies."""
    if spec.filename.endswith(".pdf"):
        return write_pdf(spec, out_dir)
    if spec.filename.endswith(".docx"):
        return write_docx(spec, out_dir)
    raise ValueError(f"unsupported extension: {spec.filename}")


def verify() -> int:
    """Check that every anchor really appears in the documents it names.

    A planted defect whose anchor is absent scores as a permanent miss and
    would quietly cap the benchmark, so this runs before generation.
    """
    by_name = {d.filename: d for d in ALL_DOCUMENTS}
    corpus_text = "\n".join(d.title + "\n" + "\n".join(d.paragraphs) for d in ALL_DOCUMENTS)
    problems: list[str] = []

    for defect in PLANTED_DEFECTS:
        for filename in defect.documents:
            if filename not in by_name:
                problems.append(f"{defect.defect_id}: missing document {filename}")

        if defect.documents:
            scope = "\n".join(
                by_name[f].title + "\n" + "\n".join(by_name[f].paragraphs)
                for f in defect.documents
                if f in by_name
            )
        else:
            scope = corpus_text

        if not any(a in scope for a in defect.anchors):
            problems.append(f"{defect.defect_id}: no anchor found in its documents")

    if problems:
        print("Consistency problems:")
        for p in problems:
            print(f"  {p}")
        return 1

    print(f"OK: {len(ALL_DOCUMENTS)} documents, {len(PLANTED_DEFECTS)} defects consistent")
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate the synthetic data room")
    parser.add_argument("--out", default="data/synthetic", type=Path)
    parser.add_argument("--verify", action="store_true", help="check consistency only")
    args = parser.parse_args()

    if args.verify:
        return verify()

    if verify() != 0:
        return 1

    out_dir: Path = args.out
    out_dir.mkdir(parents=True, exist_ok=True)

    written: list[str] = []
    for category, specs in CATEGORIES.items():
        print(f"\n  {category}")
        for spec in specs:
            path = render(spec, out_dir)
            written.append(path.name)
            print(f"    {path.name}")

    by_class: dict[str, int] = {}
    by_difficulty: dict[str, int] = {}
    for defect in PLANTED_DEFECTS:
        cls = primary_class(defect)
        by_class[cls] = by_class.get(cls, 0) + 1
        by_difficulty[defect.difficulty] = by_difficulty.get(defect.difficulty, 0) + 1

    truth_path = out_dir / "ground_truth.json"
    truth_path.write_text(
        json.dumps(
            {
                "documents": written,
                "defects": [d._asdict() for d in PLANTED_DEFECTS],
                "by_class": by_class,
                "by_difficulty": by_difficulty,
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    print(f"\n  ground_truth.json")
    print(f"\n{len(written)} documents, {len(PLANTED_DEFECTS)} planted defects")
    print("  by class:")
    for cls, n in sorted(by_class.items()):
        print(f"    {cls:<28} {n}")
    print("  by difficulty:")
    for level in ("easy", "medium", "hard"):
        print(f"    {level:<28} {by_difficulty.get(level, 0)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

"""Format-specific text extraction.

Both parsers return the same shape: full text plus the character offset at
which each page begins. Offsets are what make span validation possible, so
they are computed here once and never recomputed downstream.
"""

from __future__ import annotations

import bisect
from pathlib import Path
from typing import NamedTuple

from loupe.models.document import ParseStatus

MIN_CHARS_FOR_TEXT_LAYER = 20
PAGE_SEPARATOR = "\n\n"


class ParsedText(NamedTuple):
    """Extraction result.

    Attributes:
        text: Full document text, pages joined by PAGE_SEPARATOR.
        page_offsets: Character index at which each page starts. Index i
            corresponds to page i+1, since pages are 1-indexed for humans.
        page_count: Number of pages.
        status: OK, or the specific reason extraction failed.
        error: Human-readable detail when status is not OK.
    """

    text: str
    page_offsets: tuple[int, ...]
    page_count: int
    status: ParseStatus
    error: str | None = None


def _empty(status: ParseStatus, error: str) -> ParsedText:
    return ParsedText("", (), 0, status, error)


def page_for_offset(offset: int, page_offsets: tuple[int, ...]) -> int:
    """Return the 1-indexed page containing a character offset.

    Args:
        offset: Character index into the full document text.
        page_offsets: Start offset of each page.

    Returns:
        1-indexed page number. Returns 1 when page_offsets is empty.
    """
    if not page_offsets:
        return 1
    return max(1, bisect.bisect_right(page_offsets, offset))


def parse_pdf(path: Path) -> ParsedText:
    """Extract text from a PDF, preserving page boundaries.

    Failures are returned as ParseStatus values rather than raised. FR-5
    requires a specific reason so an unreadable document becomes an entry in
    the gap report rather than a silent hole.
    """
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except PdfReadError as exc:
        return _empty(ParseStatus.CORRUPT, str(exc))
    except Exception as exc:  # noqa: BLE001 - third-party raises broadly
        return _empty(ParseStatus.CORRUPT, f"{type(exc).__name__}: {exc}")

    if reader.is_encrypted:
        try:
            if reader.decrypt("") == 0:
                return _empty(ParseStatus.ENCRYPTED, "password required")
        except Exception:  # noqa: BLE001
            return _empty(ParseStatus.ENCRYPTED, "password required")

    chunks: list[str] = []
    offsets: list[int] = []
    cursor = 0
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            page_text = ""
        offsets.append(cursor)
        chunks.append(page_text)
        cursor += len(page_text) + len(PAGE_SEPARATOR)

    text = PAGE_SEPARATOR.join(chunks)

    if len(text.strip()) < MIN_CHARS_FOR_TEXT_LAYER:
        return _empty(
            ParseStatus.NO_TEXT_LAYER,
            "fewer than 20 extractable characters; likely a scanned image",
        )

    return ParsedText(text, tuple(offsets), len(reader.pages), ParseStatus.OK)


def parse_docx(path: Path) -> ParsedText:
    """Extract text from a DOCX file.

    DOCX has no reliable page concept without rendering, so the whole
    document is treated as page 1. This is honest rather than guessed.
    """
    from docx import Document as DocxDocument

    try:
        docx = DocxDocument(str(path))
    except Exception as exc:  # noqa: BLE001
        return _empty(ParseStatus.CORRUPT, f"{type(exc).__name__}: {exc}")

    paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
    for table in docx.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n\n".join(paragraphs)

    if len(text.strip()) < MIN_CHARS_FOR_TEXT_LAYER:
        return _empty(ParseStatus.NO_TEXT_LAYER, "document contains no text")

    return ParsedText(text, (0,), 1, ParseStatus.OK)